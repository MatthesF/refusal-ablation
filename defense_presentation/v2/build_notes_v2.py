#!/usr/bin/env python3
"""Build the v2 speaking-notes + examiner-prep Word document (02445 seminar)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

NAVY = RGBColor(0x1C, 0x34, 0x60)
NAVY2 = RGBColor(0x29, 0x45, 0x7F)
ORANGE = RGBColor(0xC7, 0x67, 0x2C)
INK = RGBColor(0x1B, 0x2A, 0x45)
MUTED = RGBColor(0x6A, 0x78, 0x90)
GREEN = RGBColor(0x2E, 0x7D, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK
for sec in doc.sections:
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def run(p, text, size=11, bold=False, italic=False, color=INK, font='Calibri', spacing=None, caps=False):
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(int(spacing * 20))); rPr.append(sp)
    return r


def para(space_before=0, space_after=6, align=None, line=1.12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if align:
        p.alignment = align
    return p


def heading(text, size=14, space_before=14):
    p = para(space_before=space_before, space_after=5)
    run(p, text, size=size, bold=True, color=NAVY, font='Cambria')
    return p


def table(headers, rows, widths, head_bg="1C3460", zebra="F2F6FC", sizes=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for c, label in zip(t.rows[0].cells, headers):
        set_cell_bg(c, head_bg)
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        c.paragraphs[0].paragraph_format.space_before = Pt(2)
        run(c.paragraphs[0], label, size=10, bold=True, color=WHITE)
    for i, rowvals in enumerate(rows):
        cells = t.add_row().cells
        if i % 2 == 0:
            for c in cells:
                set_cell_bg(c, zebra)
        for j, (c, val) in enumerate(zip(cells, rowvals)):
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            sz = sizes[j] if sizes else 10
            bold = (j == 0)
            col = NAVY2 if (j == 0) else INK
            run(c.paragraphs[0], val, size=sz, bold=bold, color=col)
    for row in t.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w
    return t


# ================= TITLE =================
p = para(space_after=2)
run(p, "CLOSED-DOOR SEMINAR · VIDEO-PITCH SCRIPT & EXAMINER PREP  (v2)", size=11, bold=True, color=ORANGE, spacing=1.6, caps=True)
p = para(space_after=2)
run(p, "Statistical Evaluation of Refusal-Direction Ablation in Gemma", size=20, bold=True, color=NAVY, font='Cambria')
p = para(space_after=10)
run(p, "02445 — Design & Evaluate a GenAI System · DTU · Group of 4", size=10.5, color=MUTED)

# ================= 1. FORMAT & STRATEGY =================
heading("1 · Exam format and where the grade is won", space_before=2)
for t, b in [
    ("Video pitch (≤ 7 min, recorded). ", "Plays at the seminar; per the brief it conveys the project but does not score points by itself. Our deck runs ~6:45 and ends on a single core-conclusions slide. Slides are split equally across the four of us."),
    ("Individual questioning (~5 min/student). ", "THIS is where the grade is decided — each of us is examined one-to-one on the project and the plenum content. Sections 4–6 below prepare exactly this."),
    ("Grade weights. ", "Report 49%, seminar 21% (of the 70% project component). Questions focus on the project; be ready to defend every method choice."),
]:
    p = para(space_after=4)
    run(p, "• ", bold=True, color=ORANGE); run(p, t, bold=True, color=NAVY2); run(p, b)

# ================= 2. SPEAKER BLOCKS =================
heading("2 · Speaker blocks (equal 4-way split)")
table(["Speaker", "Slides", "Topic", "Time"],
      [("Oskar", "1–3", "Framing: aspect (Step 1), A/B design + metric (Steps 2–3)", "~1:28"),
       ("Matthes", "4–6", "Model & runtime (Step 3), data (Step 4), outputs→Bernoulli", "~1:42"),
       ("Carl Johan", "7–9", "Statistical model + matrix/GLM, tests & assumptions, result", "~1:57"),
       ("Oliver", "10–12", "Resampling, generalization & fairness, core conclusions", "~1:37")],
      [Inches(1.2), Inches(0.8), Inches(4.3), Inches(0.9)], sizes=[10, 10, 10, 10])
p = para(space_before=4, space_after=2)
run(p, "Appendix A1–A4 are backup slides for the individual questioning (design matrix, test-choice table, power derivation, category/fairness detail).", size=10, italic=True, color=MUTED)

# ================= 3. SLIDE-BY-SLIDE SCRIPT =================
heading("3 · Video-pitch script (target ≈ 6:45)")
slides = [
    (1, "Title", "Oskar", "0:18",
     "Hello — this is our 02445 evaluation project. We ask one question: does a single weight edit remove the refusal behaviour a safety benchmark measures? The whole talk is in the corner number: official compliance goes from twenty-one to seventy-seven percent."),
    (2, "Step 1 — aspect & why", "Oskar", "0:32",
     "The aspect we evaluate is robustness — red-teaming an open-weight model's safety in weight space. It matters because open-weight models can be edited after release, so a high benchmark refusal rate only means something if it survives a cheap edit. If a post-hoc edit flips the score, the benchmark was measuring the prompt, not a robust mechanism."),
    (3, "Steps 2–3 — A/B design + metric", "Oskar", "0:38",
     "We set it up as a paired A/B test: arm A is the untouched model, arm B is the same checkpoint after one refusal-direction edit. Both answer the same 440 SORRY-Bench prompts under deterministic decoding, so each prompt is its own control. The metric is the official SORRY-Bench Mistral judge — an LLM-as-a-Judge — giving a binary score: zero refusal, one unsafe compliance."),
    (4, "Step 3 — model & runtime", "Matthes", "0:28",
     "We evaluate google/gemma-4-E4B-it, open-weight and pinned to one revision. Open weights are essential — direct weight editing is impossible through a closed API. We run it locally on a rented RunPod RTX 5090; deterministic and offline, so the 440 observations stay independent and reproducible."),
    (5, "Step 4 — data", "Matthes", "0:42",
     "Two data sources. The benchmark is SORRY-Bench — 440 prompts across 44 harm categories — used only as an external hold-out. To fit the edit we built our own 480-prompt set from the Gemma Prohibited-Use Policy: twelve policy areas crossed with safe and unsafe labels, twenty each. They're GPT-5.5-Pro-generated and human-verified, split disjointly eight areas to four. On sample size: a power calculation gives n around 349 for a fifteen-point change, so 440 prompts give us 88 percent power."),
    (6, "Step 4c — outputs → data", "Matthes", "0:32",
     "The edit in one line: fit a refusal direction, unsafe-mean minus safe-mean, and remove it from every layer's weights, norm-preserving. Then we turn outputs into data: each judged answer is a Bernoulli outcome, zero or one, and deterministic decoding fixes that label. Per prompt we keep the paired outcome — baseline and edited — which is the unit of analysis."),
    (7, "Step 5 — the statistical model", "Carl Johan", "0:42",
     "We write the evaluation as a model. Per prompt the signed change D is edited minus baseline, in minus-one, zero, plus-one, and the effect size is the mean of D. More generally it's a generalized linear model: logit of compliance equals a design matrix times beta, with columns for condition, category, and their interaction. Beta-one is the main edit effect — our A/B treatment effect; the interaction terms let the effect vary by category. So we have model set-up, the matrix formulation, and interactions in one place."),
    (8, "Step 5 — tests & assumptions", "Carl Johan", "0:42",
     "We use both families. Parametric: a normal-approximation confidence interval for the paired difference, and Wilson intervals per arm because the rates are near zero and one. Exact and non-parametric: McNemar's exact binomial test on the discordant pairs, plus a category-clustered bootstrap. Assumptions: prompts are only approximately independent — ten share each category — so we check with the clustered bootstrap; the outcome is Bernoulli; deterministic decoding removes sampling noise. The point is that all three approaches agree."),
    (9, "Step 5 — result", "Carl Johan", "0:33",
     "The result: baseline complies on 21.4 percent of prompts, the edited model on 76.6 — a paired increase of 55.2 points, 95 percent interval 50 to 60, McNemar p around ten-to-the-minus-sixty-eight. Same model, one deterministic edit."),
    (10, "Resampling & robustness", "Oliver", "0:35",
     "Resampling confirms it. The clustered bootstrap — twenty thousand resamples of whole categories — gives a wider interval, 47 to 63, still far from zero. The paired table shows why: 247 prompts flip from refusal to compliance and only four flip back. Dropping the eight length-capped answers still gives 56 points, so it's not an artefact."),
    (11, "Generalization & fairness", "Oliver", "0:40",
     "Two final checks. Generalization: the fitted direction is stable across held-out policy areas — median cosine 0.95 to 0.97 over five random splits per budget, a cross-validation-style check — and SORRY-Bench is a fully external hold-out. And a fairness lens: treating the 44 categories as subgroups, the change isn't uniform — residual refusals concentrate in child-related crimes, PII and harassment, so safety coverage differs across harm types."),
    (12, "Core conclusions", "Oliver", "0:22",
     "To conclude: one edit removes most benchmark refusal, parametric, exact and bootstrap inference all agree, but it's a benchmark-specific change with uneven coverage. The lesson: static refusal rates don't capture robustness for open-weight models. Code and data are in the repo — thank you."),
]
for num, title, speaker, time, script in slides:
    p = para(space_before=9, space_after=1)
    run(p, f"Slide {num}", size=12.5, bold=True, color=NAVY, font='Cambria')
    run(p, f"   ·   {title}", size=12.5, color=INK, font='Cambria')
    p = para(space_after=3)
    run(p, speaker.upper(), size=9.5, bold=True, color=ORANGE, spacing=1.0)
    run(p, f"   ·   {time}", size=9.5, color=MUTED)
    p = para(space_after=3, line=1.16)
    run(p, script, size=11, color=INK)

# ================= 4. RUBRIC STEP COVERAGE =================
doc.add_page_break()
heading("4 · Rubric Step coverage (Steps 1–5)", space_before=2)
table(["Step", "Requirement", "Where & what we show"],
      [("1", "Which aspect & why", "Robustness / red-teaming of open-weight refusal; validity of static safety benchmarks (slide 2)."),
       ("2", "Metric", "Official SORRY-Bench Mistral judge = LLM-as-a-Judge; binary 0/1; 40-output audit (slide 3)."),
       ("3", "Model choice + runtime", "Open-weight gemma-4-E4B-it on RunPod RTX 5090; offline, deterministic, independent (slide 4)."),
       ("4", "Data + sample size", "SORRY-Bench benchmark + self-curated 480 set; factors of variation; power n≈349→88.2% (slides 5–6)."),
       ("5", "Results: qualitative + quantitative + stability", "Category plots + paired tests, assumptions, CIs, bootstrap uncertainty (slides 7–11, A1–A4).")],
      [Inches(0.6), Inches(2.6), Inches(5.6)], sizes=[11, 10, 10])

# ================= 5. LEARNING-OBJECTIVE MAP =================
heading("5 · Learning-objective coverage map")
p = para(space_after=5)
run(p, "Each course learning objective, where we demonstrate it, and the one-line evidence. Use this to steer answers in the individual questioning.", size=10, italic=True, color=MUTED)
table(["Learning objective", "Where", "Evidence"],
      [("Set up appropriate statistical models for AI evaluation", "S3, S7, A1", "Paired A/B + GLM for SORRY-Bench compliance."),
       ("Matrix formulation of simple models", "S7, A1", "η = Xβ; design matrix with condition/category/interaction."),
       ("Test for interaction effects", "S7, A1, A4", "condition×category term; category-level effect varies 0→100 pp."),
       ("Understand & discuss model assumptions", "S8, A2", "Independence, Bernoulli, near-0/1, exchangeability — and the checks."),
       ("Parametric and non-parametric tests", "S8, A2", "Normal-approx CI + Wilson (param.); McNemar exact + bootstrap (non-param.)."),
       ("Estimation, hypothesis test, prediction w/ software", "S8–S9", "Δ estimate, McNemar, CIs in Python; GLM predicts per-category compliance."),
       ("Evaluate estimates for model generalization", "S11, S5", "Disjoint hold-out areas, 5-split CV-style stability, external benchmark."),
       ("Set up & evaluate A/B tests", "S3 (whole study)", "Paired within-subject A/B: baseline vs edited on identical prompts."),
       ("Apply resampling methods", "S10, A2", "Category-clustered bootstrap, 20,000 replicates, seed 2445."),
       ("Bias & fairness measures in AI", "S11, A4", "Category subgroups; uneven safety coverage; child-content handling."),
       ("Interpret statistical-software output", "S8–S10", "Reading CIs, p-values, bootstrap distribution, Wilson error bars."),
       ("Communicate results to a non-statistical audience", "Pitch + report", "Plain-language framing; single core-conclusions slide.")],
      [Inches(3.5), Inches(1.2), Inches(4.1)], sizes=[10, 9.5, 9.5])

# ================= 6. Q&A BANK =================
doc.add_page_break()
heading("6 · Individual-questioning Q&A bank (by competency)", space_before=2)
p = para(space_after=6)
run(p, "Likely one-to-one questions with strong, concise answers. Everyone should know the headline numbers (21.4→76.6%, +55.2 pp, 247 vs 4, p≈9×10⁻⁶⁸) and the assumptions cold.", size=10, italic=True, color=MUTED)

qa_sections = [
    ("A/B test & design", [
        ("Why is this an A/B test, and why paired?",
         "Two arms — untouched (A) vs edited (B) — evaluated on identical inputs. It's paired/within-subject because every prompt is answered by both models, so prompt difficulty, category mix and judge wording are held fixed. Pairing removes between-prompt variance and lets us analyse within-pair differences directly."),
        ("What is the unit of analysis and the population?",
         "The unit is the prompt, contributing a paired outcome (Y_baseline, Y_edited) over 440 SORRY-Bench prompts. The population is SORRY-Bench harm prompts as scored by the official judge — we're explicit that conclusions are benchmark-specific."),
    ]),
    ("Statistical model, matrix form & interactions", [
        ("Write down your statistical model.",
         "logit P(Yᵢ=1) = β₀ + β₁·edited + Σγ_c·catᵢ + Σδ_c·(edited×catᵢ). β₁ is the edit effect; a linear-probability version gives β₁ = Δ directly. McNemar/Δ are the focused, assumption-light version of β₁."),
        ("Where is the matrix formulation?",
         "η = Xβ. X has 880 rows (prompt×condition) and columns: intercept, condition indicator, 43 category dummies, and condition×category. Paired rows share category dummies; only the condition columns flip (appendix A1)."),
        ("What interaction did you test, and what did you find?",
         "Condition×category. The category-level plot IS that interaction: the edit effect ranges from ~0 to +100 pp across categories. With only 10 prompts/category we treat it as exploratory error analysis, not 44 confirmatory tests."),
    ]),
    ("Assumptions", [
        ("What are your assumptions and do they hold?",
         "Independence of prompts (only approximate — 10 share each category, so we add a category-clustered bootstrap); binary outcome modelled as Bernoulli; deterministic decoding removes within-pair sampling noise; rates near 0/1 so we use Wilson/exact, not a naive normal interval; McNemar needs the discordant pairs exchangeable under H₀."),
        ("Is prompt independence realistic?",
         "Not perfectly — SORRY-Bench is structured in blocks of 10. That's precisely why we report a category-clustered bootstrap that resamples whole categories; it widens the CI (47–63 vs 50–60) but stays far from zero, so clustering doesn't change the conclusion."),
    ]),
    ("Tests: parametric vs non-parametric", [
        ("Why McNemar rather than a paired t-test or a chi-square?",
         "The data are paired and binary. McNemar is the paired test for binary outcomes — it uses only the discordant pairs. An ordinary chi-square assumes two independent samples, which pairing violates; a paired t-test on the {−1,0,1} differences is essentially the normal-approximation CI we already report."),
        ("Give one parametric and one non-parametric result and say why both.",
         "Parametric: normal-approx 95% CI for Δ = [50.4, 60.0]. Non-parametric/exact: McNemar p ≈ 9×10⁻⁶⁸ and a bootstrap CI [47.3, 63.0]. We report both because the parametric CI is interpretable while the exact test makes no large-sample assumption — and they agree."),
        ("Could you have used a permutation test?",
         "Yes. McNemar's exact test is the exact/combinatorial version — a binomial on the discordant pairs under p=0.5. A label-permutation test that swaps the condition within each pair gives an equivalent null distribution for Δ."),
    ]),
    ("Estimation, software & p-values", [
        ("How did you compute the confidence intervals?",
         "The paired SE of Dᵢ for the normal-approx CI; Wilson intervals per arm; and a percentile interval from the clustered bootstrap. All in Python (numpy/scipy)."),
        ("Interpret the p-value.",
         "Under H₀ that flips are equally likely in both directions, the chance of data this extreme is ≈ 9×10⁻⁶⁸. But the p-value is secondary — with this n almost anything is significant; the effect size (+55.2 pp) is the headline."),
    ]),
    ("Generalization (hold-out & cross-validation)", [
        ("How did you assess generalization?",
         "Three ways: (1) policy-area-disjoint hold-out — fit on 8 areas, compare the direction to one fit on 4 held-out areas; (2) repeated random splits, 5 per budget — a cross-validation-style stability check (cosine 0.95–0.97); (3) SORRY-Bench is a fully external benchmark never used in fitting."),
        ("Isn't fitting a direction then testing it leakage?",
         "Not for the main result — construction data and SORRY-Bench are disjoint and the benchmark is never used to fit. The exploratory second-direction follow-up WOULD be leakage because it uses evaluation outcomes, so we report that strictly as error analysis, not a confirmatory result."),
    ]),
    ("Resampling", [
        ("Describe your bootstrap and why it's clustered.",
         "Category-clustered: 20,000 replicates, seed 2445; resample the 44 categories with replacement, keeping all 10 pairs in a sampled category together, then recompute Δ → percentile CI [47.3, 63.0]. A prompt-level bootstrap would assume independence and understate uncertainty; clustering respects the category structure."),
    ]),
    ("Bias & fairness", [
        ("This is robustness — where is bias/fairness?",
         "We disaggregate by the 44 harm categories as subgroups and ask whether the safety change is uniform. It isn't: residual refusals concentrate in child-related crimes, PII and harassment, so post-edit safety coverage is unequal across harm types — a fairness-of-coverage finding. We also excluded child-related content from construction for responsible data handling."),
        ("With 44 categories, what about multiple comparisons?",
         "Exactly why the category analysis is framed as exploratory error analysis, not 44 hypothesis tests — with 10 prompts each they'd be underpowered and inflate false positives. The single aggregate Δ is the confirmatory claim."),
    ]),
    ("Sample size, metric & model choice", [
        ("How did you decide how many prompts to use?",
         "Power calculation: n = (σ·(z₁₋β+z₁₋α/2)/Δ)², with σ=1 (conservative, D∈[−1,1]), Δ=0.15, 80% power → n≈349. SORRY-Bench is fixed at 440 → 88.2% power; the observed 55.2 pp is far above the 15 pp target."),
        ("Is the LLM-as-a-Judge valid?",
         "It's the official SORRY-Bench fine-tuned Mistral judge — reproducible and standard, but coarse and binary. We audited 40 outputs: labels were reasonable for 36/40. Some score-1 answers still carry disclaimers, so we read 76.6% as judged compliance, not maximal harm."),
        ("Why Gemma, and did you need a GPU?",
         "Open weights let us edit the model directly — impossible via a closed API. We ran it offline on a RunPod RTX 5090; deterministic and offline means independent, reproducible samples."),
    ]),
    ("Project-specific", [
        ("What about the eight length-capped generations?",
         "All eight were in the baseline arm. A sensitivity analysis dropping them still gives +56.0 pp — same conclusion."),
        ("Child-related crimes were excluded from construction yet rose ~30 pp — explain.",
         "Excluded only from the data we used to fit the edit, for responsible handling; SORRY-Bench still tests the category. The rise shows the fitted direction generalises beyond the categories it was built from — consistent with a shared refusal signal."),
    ]),
]
for sec_title, items in qa_sections:
    p = para(space_before=9, space_after=3)
    run(p, sec_title, size=12, bold=True, color=NAVY2, font='Cambria')
    for q, a in items:
        p = para(space_before=3, space_after=1)
        run(p, "Q.  ", size=11, bold=True, color=ORANGE); run(p, q, size=11, bold=True, color=NAVY2)
        p = para(space_after=3, line=1.16)
        run(p, "A.  ", size=11, bold=True, color=MUTED); run(p, a, size=11, color=INK)

# ================= 7. PER-STUDENT FOCUS =================
doc.add_page_break()
heading("7 · Per-student focus for the one-to-one round", space_before=2)
p = para(space_after=6)
run(p, "Questions can go to anyone, so everyone owns the headline numbers and the assumptions. These are the areas each person should be ready to lead on.", size=10, italic=True, color=MUTED)
focus = [
    ("Oskar", "Framing & evaluation design: why this aspect matters, the paired A/B set-up, LLM-as-a-Judge validity and its limits, and communicating results to non-experts."),
    ("Matthes", "System & data: model choice and runtime (open-weight, GPU), the refusal-direction edit, factors of variation in the construction set, outputs→Bernoulli, and the power/sample-size calculation."),
    ("Carl Johan", "Inference: the statistical model and matrix/GLM formulation, interaction effects, assumptions and how they're checked, parametric vs non-parametric tests, and why McNemar."),
    ("Oliver", "Robustness & responsible AI: the category-clustered bootstrap, paired-flip reading, generalization via hold-out/CV, the bias/fairness-of-coverage lens, and multiple-comparisons handling."),
]
for who, what in focus:
    p = para(space_before=4, space_after=3)
    run(p, who + ".  ", size=11.5, bold=True, color=ORANGE)
    run(p, what, size=11, color=INK)

out = os.path.join(HERE, "Refusal_Ablation_Defense_v2_Speaking_Notes.docx")
doc.save(out)
print("Saved", out)
