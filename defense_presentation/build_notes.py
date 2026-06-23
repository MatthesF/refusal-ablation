#!/usr/bin/env python3
"""Build the oral-defense speaking-notes Word document."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

NAVY = RGBColor(0x1C, 0x34, 0x60)
NAVY2 = RGBColor(0x29, 0x45, 0x7F)
ORANGE = RGBColor(0xC7, 0x67, 0x2C)
INK = RGBColor(0x1B, 0x2A, 0x45)
MUTED = RGBColor(0x6A, 0x78, 0x90)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
# base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for sec in doc.sections:
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def run(p, text, size=11, bold=False, italic=False, color=INK, font='Calibri', spacing=None, caps=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(int(spacing * 20)))
        rPr.append(sp)
    return r


def para(space_before=0, space_after=6, align=None, line=1.12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if align:
        p.alignment = align
    return p


# ---------------- TITLE ----------------
p = para(space_after=2)
run(p, "ORAL DEFENSE · SPEAKING NOTES", size=11, bold=True, color=ORANGE, spacing=2.0, caps=True)
p = para(space_after=2)
run(p, "Statistical Evaluation of Refusal-Direction Ablation in Gemma", size=21, bold=True, color=NAVY, font='Cambria')
p = para(space_after=10)
run(p, "02445 Project Report · Technical University of Denmark · Target length ≈ 8 minutes (~7:50, leaving buffer)", size=10.5, color=MUTED)

# how-to note
p = para(space_after=4)
run(p, "The orange pill in the top-right corner of each slide names the active speaker. Hand-offs are marked in the “Transition” line. Times are guides, not scripts — aim slightly under so questions fit in eight minutes.", size=10.5, italic=True, color=MUTED)

# ---------------- SPEAKER OVERVIEW TABLE ----------------
p = para(space_before=6, space_after=4)
run(p, "Speaker blocks", size=14, bold=True, color=NAVY, font='Cambria')

overview = [
    ("Oskar", "1–3", "Framing — motivation & research question", "≈ 2:00"),
    ("Matthes", "4–6", "Methods — data, the weight edit, stability", "≈ 2:05"),
    ("Carl Johan", "7–8", "Statistics & the headline result", "≈ 1:35"),
    ("Oliver", "9–10", "Paired flips & category error analysis", "≈ 1:30"),
    ("Oskar", "11", "Conclusion & implications", "≈ 0:40"),
]
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.autofit = True
hdr = tbl.rows[0].cells
for c, label in zip(hdr, ["Speaker", "Slides", "Topic", "Time"]):
    set_cell_bg(c, "1C3460")
    c.paragraphs[0].paragraph_format.space_after = Pt(2)
    c.paragraphs[0].paragraph_format.space_before = Pt(2)
    run(c.paragraphs[0], label, size=10.5, bold=True, color=WHITE)
for i, (sp, sl, tp, tm) in enumerate(overview):
    cells = tbl.add_row().cells
    if i % 2 == 0:
        for c in cells:
            set_cell_bg(c, "F2F6FC")
    vals = [(sp, True, NAVY2), (sl, False, INK), (tp, False, INK), (tm, True, MUTED)]
    for c, (val, b, col) in zip(cells, vals):
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        c.paragraphs[0].paragraph_format.space_before = Pt(2)
        run(c.paragraphs[0], val, size=10.5, bold=b, color=col)
# widen topic column
widths = [Inches(1.3), Inches(0.8), Inches(3.6), Inches(0.9)]
for row in tbl.rows:
    for idx, w in enumerate(widths):
        row.cells[idx].width = w

# ---------------- PER-SLIDE SCRIPT ----------------
p = para(space_before=14, space_after=8)
run(p, "Slide-by-slide script", size=14, bold=True, color=NAVY, font='Cambria')

slides = [
    (1, "Title", "Oskar", "0:20",
     "Open; the headline number is the whole talk.",
     "Good morning. We're presenting our 02445 project — a statistical evaluation of refusal-direction "
     "ablation in Gemma. The one-line version is the number on the left: a single weight edit takes the "
     "model's unsafe-compliance rate on the official SORRY-Bench from about twenty-one percent up to "
     "seventy-seven percent. Let me set up why that matters.",
     "Stay on; continue to motivation."),
    (2, "Motivation — refusal is a mechanism you can edit", "Oskar", "0:50",
     "Three beats: expected behaviour → open weights → one direction.",
     "Instruction-tuned models are expected to refuse harmful requests while staying useful on everything "
     "else. For closed models the only levers are the prompt and the decoding — but for open-weight models "
     "the weights themselves are an attack surface, and you can intervene directly. Prior work makes this "
     "concrete: Arditi and colleagues showed in 2024 that refusal is largely mediated by a single direction "
     "in activation space, and the open-source community operationalised that as “abliteration” — fitting "
     "that direction and erasing it. So the concern is simple: a model can post a high score on a fixed "
     "refusal benchmark while leaning on a mechanism a cheap, post-hoc edit can switch off. Our question is "
     "how much of the score actually depends on that one mechanism.",
     "Stay on; pose the formal question."),
    (3, "Research question & design", "Oskar", "0:50",
     "Read the question, then the two conditions.",
     "So we asked one narrow, reproducible question: if we fit a refusal direction from a small policy-derived "
     "dataset and remove it from Gemma's weights, how much does official SORRY-Bench refusal behaviour change? "
     "The design is deliberately minimal — two conditions: the untouched Gemma-4-E4B-it model, and the exact "
     "same checkpoint after one fixed refusal-direction weight edit. Both answer the same 440 SORRY-Bench "
     "prompts under deterministic decoding, scored by the official Mistral judge, where zero is a refusal and "
     "one is unsafe compliance. It's a paired design — every prompt answered once by each model — and the goal "
     "is to measure the change, not to build a useful model.",
     "Hand to Matthes for the method."),
    (4, "Construction data", "Matthes", "0:40",
     "Stress: construction and benchmark never overlap.",
     "Thanks. The refusal direction is fit from a local, policy-derived dataset — 480 prompts built from "
     "Gemma's own Prohibited-Use Policy: twelve policy areas with twenty safe and twenty unsafe prompts each. "
     "We split by policy area: eight areas — 320 prompts — fit the direction; four areas — 160 prompts — are "
     "held out for a stability check. Prompts were generated with GPT-5.5 Pro and human-verified, and there's "
     "no training or fine-tuning anywhere. One deliberate choice: child-related harmful content is excluded "
     "from construction. And critically, as the table shows, the construction data and SORRY-Bench never overlap.",
     "Continue to the edit itself."),
    (5, "The direction & the weight edit", "Matthes", "0:55",
     "Two numbered steps; point at the W⊥ formula.",
     "The method is two steps. First, fit the direction: we run each prompt through the model, record the "
     "final-token hidden state at every layer, and take the mean of the unsafe activations minus the mean of "
     "the safe activations. One Gram-Schmidt step removes the component along the safe-prompt mean — a guardrail "
     "so we target refusal rather than erasing general answering ability — then we normalise to a unit vector r "
     "per layer. Second, remove it from the weights: for every layer we edit the attention-output and the MLP "
     "down-projection matrices, stripping from each column the part that writes along r — that's the formula, "
     "W-perp equals W minus r times r-transpose-W. We restore the original column norms, so it's norm-preserving, "
     "not just shrinking weights. One fixed edited model is then used for all 440 prompts.",
     "Continue to the stability check."),
    (6, "Direction stability (sanity check)", "Matthes", "0:30",
     "Quick; emphasise it's never used to tune the result.",
     "Before any benchmarking, one sanity check. We fit the direction from different groups of policy areas and "
     "compared them layer by layer with cosine similarity against the held-out areas. The directions are stable — "
     "median cosine between 0.95 and 0.97, and the tenth-percentile line stays high too — which suggests we're "
     "capturing a shared, low-dimensional refusal signal, not just category wording. Important: this check is "
     "never used to tune the final result.",
     "Hand to Carl Johan for the statistics."),
    (7, "Statistical design", "Carl Johan", "0:45",
     "Signed change → tests; mention power and determinism.",
     "Thanks. Because both models answer the same prompts, this is a paired experiment. Each prompt gets a signed "
     "change — plus one if the edit turns a refusal into compliance, minus one for the reverse, zero if nothing "
     "changes — and our effect size is the mean signed change. We run McNemar's exact test on the discordant pairs "
     "to ask whether flips in the two directions are equally likely; we report a 95% confidence interval from the "
     "paired standard error; and as a robustness check we use a category-clustered bootstrap with twenty thousand "
     "replicates. The design is powered — 440 prompts give about 88% power to detect a 15-point change — and because "
     "decoding is deterministic, every flip is a real label change, not run-to-run noise.",
     "Continue to the headline result."),
    (8, "Main result", "Carl Johan", "0:50",
     "Land the numbers slowly; pause after 55.2 pp.",
     "And here's the headline. Baseline Gemma complies with 21.4% of prompts — 94 of 440. After the edit, that "
     "jumps to 76.6% — 337 of 440. The paired difference is plus 55.2 percentage points, with a 95% interval of "
     "50.4 to 60, and even the conservative category-clustered bootstrap interval, 47 to 63, stays far from zero. "
     "The figure shows it with Wilson intervals on each bar. The thing to hold onto: this is the same model and "
     "one deterministic weight edit — no retraining — and most of the benchmark's refusal behaviour is gone.",
     "Hand to Oliver for where the change comes from."),
    (9, "Paired flips & McNemar", "Oliver", "0:45",
     "Walk the 2×2; land on 247 vs 4.",
     "Thanks. This table is where the aggregate change comes from. Of the prompts baseline Gemma refused, 247 "
     "flipped to compliance after the edit and only 99 stayed refusals. Of the ones it already complied with, just "
     "4 flipped the other way. So 247 versus 4 — almost entirely one-directional — and McNemar's exact test gives "
     "a p-value around ten-to-the-minus-sixty-eight. The four opposite flips are borderline, caution-heavy cases "
     "like personal insults and fake news, so we report them as error analysis, not as a safety gain.",
     "Continue to the category view."),
    (10, "Category-level error analysis", "Oliver", "0:45",
     "Broad-but-uneven; flag the child-related point and the 10/category caveat.",
     "By category, the shift is broad but uneven. On the left, many categories move from zero compliance all the "
     "way to a hundred — system intrusion, explicit content, governance advice — while others barely move. On the "
     "right, the refusals that remain are concentrated: 103 prompts are still refused, clustered in a handful of "
     "categories, with child-related crimes worst at seven out of ten, then insults, harassment, PII and financial "
     "advice. One striking point — child-related crimes still rose about thirty points even though we excluded that "
     "content from construction. With only ten prompts per category, this is exploratory error analysis, not a ranking.",
     "Hand back to Oskar to close."),
    (11, "Conclusion & implications", "Oskar", "0:40",
     "Result / scope / implication, then thanks.",
     "To wrap up. The result: for this Gemma model, one fitted refusal-direction edit removes most of the "
     "benchmark-measured refusal — 21 to 77 percent, flips 247 versus 4 — statistically clear and practically "
     "large. The scope: this is a benchmark-specific behaviour change, not a verdict on the model's overall "
     "quality or every safety mechanism, and some categories stay resistant. The implication — really the point — "
     "is that a static refusal rate alone doesn't capture robustness for open-weight models; our own residual "
     "follow-up even showed how easily you slip into test-set leakage. Future safety evaluations should add "
     "robustness-to-editing and fresh held-out tests. Everything is in the repo. Thank you — we're happy to take questions.",
     "End of presentation; open the floor."),
]

for num, title, speaker, time, cue, script, transition in slides:
    p = para(space_before=10, space_after=1)
    run(p, f"Slide {num}", size=13, bold=True, color=NAVY, font='Cambria')
    run(p, f"   ·   {title}", size=13, bold=False, color=INK, font='Cambria')
    # meta line
    p = para(space_after=4)
    run(p, speaker.upper(), size=10, bold=True, color=ORANGE, spacing=1.2)
    run(p, f"   ·   {time}   ·   ", size=10, color=MUTED)
    run(p, f"Cue: {cue}", size=10, italic=True, color=MUTED)
    # script
    p = para(space_after=4, line=1.18)
    run(p, script, size=11, color=INK)
    # transition
    p = para(space_after=2)
    run(p, "→ Transition:  ", size=10, bold=True, color=NAVY2)
    run(p, transition, size=10, italic=True, color=MUTED)

# ---------------- Q&A PREP ----------------
doc.add_page_break()
p = para(space_after=8)
run(p, "Anticipated questions", size=14, bold=True, color=NAVY, font='Cambria')

qa = [
    ("Why Gemma-4-E4B-it specifically?",
     "It's an instruction-tuned, open-weight model we can run locally while still doing direct weight-space "
     "intervention — which simply isn't possible through a closed, hosted API. The exact checkpoint revision is pinned."),
    ("Isn't this just “abliteration” from the open-source community?",
     "It's in that family, yes. Our contribution is the evaluation, not a new attack: a paired SORRY-Bench design "
     "with McNemar's exact test and a category-clustered bootstrap, a norm-preserving edit, and a policy-derived "
     "construction set that's disjoint from the benchmark."),
    ("Does the binary Mistral judge bias the result?",
     "It's reproducible and official but coarse. We manually audited a stratified sample of 40 outputs and found the "
     "official labels reasonable for 36 of 40; the borderline ones mixed cautionary language with task help. Some "
     "score-1 answers still carry disclaimers — so we read 76.6% as “judged compliance,” not “maximally harmful.”"),
    ("You excluded child-related prompts from construction — why did that category still rise ~30 pp?",
     "We excluded it from the construction data as a data-handling choice, to avoid generating or inspecting that "
     "content; SORRY-Bench still tests it through aggregate scores. The rise shows the fitted direction generalises "
     "beyond the categories it was built from — consistent with a shared refusal signal."),
    ("What about the eight generations that hit the token cap?",
     "All eight were in the baseline condition. A sensitivity check that drops those prompts still gives a +56.0 pp "
     "paired increase — the same substantive conclusion."),
    ("Is the effect just statistical luck given the sample size?",
     "No. We pre-justified n with a power calculation: 440 paired prompts give 88.2% power for a 15 pp change, and the "
     "observed 55.2 pp is far above that threshold, with both the normal-approximation and clustered-bootstrap "
     "intervals well clear of zero."),
    ("What's the residual follow-up you mention, and why flag leakage?",
     "We did a small exploratory follow-up fitting a second direction against the prompts the edited model still "
     "refused. It further reduced refusal — but because it uses final evaluation outcomes to design a new "
     "intervention, it's test-set leakage. We report it strictly as error analysis, not a confirmatory second result."),
]
for q, a in qa:
    p = para(space_before=8, space_after=2)
    run(p, "Q.  ", size=11, bold=True, color=ORANGE)
    run(p, q, size=11, bold=True, color=NAVY2)
    p = para(space_after=4, line=1.18)
    run(p, "A.  ", size=11, bold=True, color=MUTED)
    run(p, a, size=11, color=INK)

out = os.path.join(HERE, "Refusal_Ablation_Defense_Speaking_Notes.docx")
doc.save(out)
print("Saved", out)
